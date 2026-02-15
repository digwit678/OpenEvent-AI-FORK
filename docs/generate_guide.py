from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('OpenEvent-AI: AI Integration Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    doc.add_paragraph('This document outlines the strategy, setup, and operational patterns for using AI within the OpenEvent-AI project.')

    # Strategy
    doc.add_heading('1. AI Strategy: Hybrid Intelligence', level=1)
    p = doc.add_paragraph()
    p.add_run('Strategy: ').bold = True
    p.add_run('The project employs a "Hybrid Intelligence" model that balances the flexibility of Large Language Models (LLMs) with the reliability of deterministic logic.')
    
    doc.add_heading('Core Principles:', level=2)
    bullets = [
        ('Semantic-First Detection', 'Using LLMs (OpenAI/Gemini) to classify intents and extract complex entities like dates and room preferences.'),
        ('Deterministic Guardrails', 'Regex and heuristic gates act as fallbacks and validators to ensure "hard facts" (prices, dates) remain 100% accurate.'),
        ('Capture-Anytime', 'Global extraction logic captures billing and contact info from any message, regardless of the current workflow step.'),
        ('Change-Anytime (Anchored)', 'A sophisticated change propagation system detects detours (e.g., changing date during billing) and anchors them to the current state.')
    ]
    for b_title, b_text in bullets:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{b_title}: ').bold = True
        p.add_run(b_text)

    # Setup
    doc.add_heading("2. Setup: The Agent's Context", level=1)
    doc.add_paragraph('To ensure the AI operates safely and efficiently, the project uses a layered context system:')
    
    setup_items = [
        ('Session Primers', 'Daily auto-generated briefs (session_primer.md) that tell the AI what changed in the last 24 hours.'),
        ('Startup Packs', 'Categorized "Deep Dive" instructions (Pack A: Routing, Pack B: Site Visit, Pack C: Verbalization) for focused development.'),
        ('Unified Instructions', 'Dedicated instruction files (AGENTS.md, GEMINI.md) defining non-negotiable architectural rules.'),
        ('Multi-LLM Hybrid Mode', 'Dynamic switching between OpenAI (for high-precision extraction) and Gemini (for fast classification and reasoning).')
    ]
    for s_title, s_text in setup_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{s_title}: ').bold = True
        p.add_run(s_text)

    # Skills & Repeated Actions
    doc.add_heading('3. Repeated Actions & Skills', level=1)
    doc.add_paragraph('The AI uses "Skills"—standardized procedures for high-risk or frequent tasks:')
    
    skills = [
        ('Architectural Guardrails', 'A pre-flight checklist used before touching the routing pipeline to prevent regression in "Lost Message" or "Wrong Step" bugs.'),
        ('Test Matrix Navigator', 'Automatically identifies the minimal set of tests needed to verify a specific change, maintaining high development velocity.'),
        ('Universal Verbalization', 'A centralized system (universal_verbalizer.py) that ensures AI-generated responses follow a specific persona while preserving data integrity.')
    ]
    for sk_title, sk_text in skills:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{sk_title}: ').bold = True
        p.add_run(sk_text)

    # Hooks
    doc.add_heading('4. Hooks & Pipelines', level=1)
    doc.add_paragraph('Automation is baked into the "Pre-Route" and "HIL" (Human-in-the-Loop) pipelines:')
    
    hooks = [
        ('Pre-Route Chain', 'A 5-stage interceptor pipeline (Duplicate check -> Out-of-Context -> Smart Shortcuts -> Billing Correction) that runs before the LLM.'),
        ('HIL Approval System', 'AI drafts for complex offers or confirmations are queued in a dedicated dashboard for manager approval before being sent to clients.'),
        ('CLI Automation Hooks', 'Shell scripts (plan-review.sh, ask-codex-debug.sh) and Apple Notes integration for external task management.')
    ]
    for h_title, h_text in hooks:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{h_title}: ').bold = True
        p.add_run(h_text)

    # Footer
    doc.add_paragraph('\nGenerated on 2026-02-11 for OpenEvent-AI Onboarding.')

    doc.save('OpenEvent_AI_Integration_Guide.docx')
    print("Document saved: OpenEvent_AI_Integration_Guide.docx")

if __name__ == "__main__":
    create_doc()
