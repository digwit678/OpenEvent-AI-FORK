from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_dev_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('OpenEvent-AI: AI Development Agent Workflow', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    doc.add_paragraph('This document explains how AI agents (Gemini/Claude) are used as active developers on the OpenEvent-AI project, focusing on the specialized hooks, skills, and programming strategies employed.')

    # 1. The Setup: Context-Driven Development
    doc.add_heading('1. The Setup: Context-Driven Development', level=1)
    doc.add_paragraph('The agent is not "blind"; it is initialized with a multi-layered context system that ensures it understands the current state of the project.')
    
    items = [
        ('Session Primer Hook', 'The agent starts every session by reading docs/daily_scrum/session_primer.md to understand recent regressions and immediate priorities.'),
        ('Base Rules (AGENTS.md)', 'A project-wide rulebook that defines non-negotiable coding standards, directory structures, and safety constraints.'),
        ('Memory Persistence', 'User-specific preferences and project-wide facts are saved to long-term memory to maintain continuity across CLI sessions.')
    ]
    for title, text in items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(text)

    # 2. Programming Strategy: Plan-Then-Execute
    doc.add_heading('2. Programming Strategy: Plan-Then-Execute', level=1)
    doc.add_paragraph('To prevent "hallucinated code" or architectural drift, the project enforces a strict workflow:')
    
    strategies = [
        ('Discovery Phase', 'The agent uses parallel grep_search and read_file calls to map out dependencies before suggesting changes.'),
        ('Explicit Planning', 'The agent must present a concise plan for approval. This is often cross-referenced with the MASTER_ARCHITECTURE_SHEET.md.'),
        ('Atomic Replacements', 'Large changes are broken down into small, targeted text replacements to minimize the risk of breaking surrounding logic.')
    ]
    for title, text in strategies:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(text)

    # 3. Agent Skills (Executable Documentation)
    doc.add_heading('3. Agent Skills: Specialized Expertises', level=1)
    doc.add_paragraph('We use the .codex/skills/ folder to give the AI specialized "Agent Skills" that can be activated on demand.')
    
    skills = [
        ('oe-architectural-guardrails', 'Activated when touching the routing pipeline. It provides a checklist of common "Bug Magnets" (e.g., date drift, OOC drops).'),
        ('oe-test-matrix-navigator', 'Identifies the minimal subset of integration tests relevant to the current change, keeping the dev-loop fast.'),
        ('oe-workflow-triage', 'A specialized skill for investigating "Lost Message" bugs by tracing state transitions in the JSON database.')
    ]
    for title, text in skills:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(text)

    # 4. Automation Hooks
    doc.add_heading('4. Automation Hooks', level=1)
    doc.add_paragraph('A set of shell-to-agent hooks bridge the gap between the terminal and the AI brain.')
    
    hooks = [
        ('plan-review.sh', 'A pre-commit hook style script that asks the agent to review a plan for architectural compliance.'),
        ('e2e-browser-checklist.sh', 'Automates browser-based verification using Playwright, providing screenshots directly to the agent.'),
        ('session-summary-hook', 'At the end of a session, the agent automatically updates the DEV_CHANGELOG.md and generates the next primer.')
    ]
    for title, text in hooks:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(text)

    doc.add_paragraph('
Final Strategy Note: The AI is treated as a "Staff Engineer" that must document its own work, verify its own fixes, and maintain the project's architectural integrity through continuous documentation updates.')

    doc.save('OpenEvent_AI_Dev_Agent_Workflow.docx')
    print("Dev Agent Document saved: OpenEvent_AI_Dev_Agent_Workflow.docx")

if __name__ == "__main__":
    create_dev_doc()
